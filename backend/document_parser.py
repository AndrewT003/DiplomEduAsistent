import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import List, Dict


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Витягує тільки текст (для зворотної сумісності)"""
    if filename.endswith(".pdf"):
        return extract_from_pdf(file_bytes)
    elif filename.endswith(".docx"):
        return extract_from_docx(file_bytes)
    else:
        raise ValueError(f"Непідтримуваний формат: {filename}")


def extract_text_with_formatting(file_bytes: bytes, filename: str) -> List[Dict]:
    """Витягує текст разом з метаданими форматування

    Returns:
        List[Dict]: Список блоків тексту з форматуванням:
        [
            {
                "text": "Текст параграфа",
                "alignment": "center|left|right|justify",
                "font_name": "Times New Roman",
                "font_size": 14,
                "bold": False,
                "italic": False,
                "page": 1  # для PDF
            }
        ]
    """
    if filename.endswith(".pdf"):
        return extract_from_pdf_with_formatting(file_bytes)
    elif filename.endswith(".docx"):
        return extract_from_docx_with_formatting(file_bytes)
    else:
        raise ValueError(f"Непідтримуваний формат: {filename}")


def extract_from_pdf(file_bytes: bytes) -> str:
    """Витягує тільки текст з PDF"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text.strip()


def extract_from_pdf_with_formatting(file_bytes: bytes) -> List[Dict]:
    """Витягує текст з PDF разом з форматуванням"""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    blocks = []

    for page_num, page in enumerate(doc, start=1):
        # Отримуємо структуровані дані про блоки тексту
        text_dict = page.get_text("dict")
        page_width = page.rect.width

        for block in text_dict.get("blocks", []):
            if block.get("type") == 0:  # Текстовий блок
                for line in block.get("lines", []):
                    line_text = ""
                    font_name = None
                    font_size = None
                    bold = False
                    italic = False

                    # Збираємо інформацію про текст та форматування
                    for span in line.get("spans", []):
                        line_text += span.get("text", "")
                        if not font_name:
                            font_name = span.get("font", "Unknown")
                        if not font_size:
                            font_size = span.get("size", 0)

                        # Визначаємо стилі за назвою шрифту
                        font_lower = span.get("font", "").lower()
                        if "bold" in font_lower:
                            bold = True
                        if "italic" in font_lower or "oblique" in font_lower:
                            italic = True

                    if line_text.strip():
                        # Визначаємо вирівнювання за положенням на сторінці
                        bbox = line.get("bbox", [0, 0, 0, 0])
                        line_x = bbox[0]
                        line_width = bbox[2] - bbox[0]
                        center_x = page_width / 2

                        # Приблизне визначення вирівнювання
                        if abs((line_x + line_width/2) - center_x) < 50:
                            alignment = "center"
                        elif line_x < 100:
                            alignment = "left"
                        elif line_x > page_width - 100:
                            alignment = "right"
                        else:
                            alignment = "left"

                        blocks.append({
                            "text": line_text.strip(),
                            "alignment": alignment,
                            "font_name": font_name,
                            "font_size": round(font_size, 1),
                            "bold": bold,
                            "italic": italic,
                            "page": page_num
                        })

    doc.close()
    return blocks


def extract_from_docx(file_bytes: bytes) -> str:
    """Витягує тільки текст з DOCX"""
    print(f"  📄 Базове витягування тексту з DOCX ({len(file_bytes)} байт)...")
    doc = Document(BytesIO(file_bytes))
    text = ""
    total_paragraphs = len(doc.paragraphs)
    print(f"  📊 Знайдено параграфів: {total_paragraphs}")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    text = text.strip()
    print(f"  ✅ Витягнуто {len(text)} символів тексту")

    if len(text) == 0:
        print(f"  ⚠️ УВАГА: Документ не містить текстового вмісту!")

    return text


def extract_from_docx_with_formatting(file_bytes: bytes) -> List[Dict]:
    """Витягує текст з DOCX разом з форматуванням"""
    print(f"  📄 Парсинг DOCX файлу ({len(file_bytes)} байт)...")
    doc = Document(BytesIO(file_bytes))
    blocks = []

    # Мапа вирівнювання
    alignment_map = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        None: "left"
    }

    total_paragraphs = len(doc.paragraphs)
    print(f"  📊 Знайдено параграфів у DOCX: {total_paragraphs}")

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            # Вирівнювання параграфа
            alignment = alignment_map.get(paragraph.alignment, "left")

            # Отримуємо форматування з першого run (частини тексту)
            font_name = "Unknown"
            font_size = 0
            bold = False
            italic = False

            if paragraph.runs:
                first_run = paragraph.runs[0]
                if first_run.font.name:
                    font_name = first_run.font.name
                if first_run.font.size:
                    # Розмір у Pt (points)
                    font_size = round(first_run.font.size.pt, 1)
                bold = first_run.font.bold or False
                italic = first_run.font.italic or False

            blocks.append({
                "text": paragraph.text.strip(),
                "alignment": alignment,
                "font_name": font_name,
                "font_size": font_size,
                "bold": bold,
                "italic": italic,
                "page": None  # DOCX не має явних сторінок
            })

    print(f"  ✅ Витягнуто {len(blocks)} блоків з непорожнім текстом")

    if len(blocks) == 0:
        print(f"  ⚠️ УВАГА: Не знайдено жодного блоку з текстом!")
        print(f"  📝 Всі {total_paragraphs} параграфів порожні або містять тільки пробіли")

    return blocks


def detect_document_structure(blocks: List[Dict]) -> List[Dict]:
    """Визначає структуру документа (розділи, підрозділи)

    Returns:
        List[Dict]: Блоки з доданим полем "section" (назва розділу)
    """
    import re

    current_section = "ПОЧАТОК ДОКУМЕНТА"
    section_patterns = [
        # РОЗДІЛ 1, РОЗДІЛ 1., Розділ 1
        r'^(?:РОЗДІЛ|Розділ|РАЗДЕЛ|Раздел|SECTION|Section)\s+(\d+)',
        # 1. Назва розділу, 1 Назва розділу
        r'^(\d+)\.?\s+[А-ЯҐЄІЇ]',
        # ВСТУП, ВИСНОВКИ і т.д.
        r'^(ВСТУП|ВИСНОВКИ|ВИСНОВОК|ЗМІСТ|СПИСОК|ДОДАТ[ОК]К|REFERENCES|ABSTRACT)',
    ]

    for block in blocks:
        text = block.get("text", "").strip()

        # Перевіряємо чи це заголовок розділу
        for pattern in section_patterns:
            match = re.match(pattern, text, re.IGNORECASE | re.UNICODE)
            if match:
                # Це початок нового розділу
                current_section = text[:100]  # Обмежуємо довжину назви
                break

        # Додаємо інформацію про поточний розділ
        block["section"] = current_section

    return blocks


def format_blocks_summary(blocks: List[Dict]) -> str:
    """Форматує інформацію про блоки для передачі в LLM

    Створює короткий опис форматування документа
    """
    if not blocks:
        return "Документ порожній"

    summary_lines = ["ІНФОРМАЦІЯ ПРО ФОРМАТУВАННЯ ТА СТРУКТУРУ ДОКУМЕНТА:\n"]

    # Групуємо по вирівнюванню
    alignment_stats = {"left": 0, "center": 0, "right": 0, "justify": 0}
    font_sizes = {}
    fonts = {}

    for block in blocks:
        alignment_stats[block.get("alignment", "left")] += 1

        font_size = block.get("font_size", 0)
        if font_size > 0:
            font_sizes[font_size] = font_sizes.get(font_size, 0) + 1

        font_name = block.get("font_name", "Unknown")
        fonts[font_name] = fonts.get(font_name, 0) + 1

    # Статистика
    summary_lines.append(f"Всього блоків тексту: {len(blocks)}")
    summary_lines.append(f"\nВирівнювання:")
    for align, count in alignment_stats.items():
        if count > 0:
            summary_lines.append(f"  - {align}: {count} блоків")

    summary_lines.append(f"\nШрифти:")
    for font, count in sorted(fonts.items(), key=lambda x: -x[1])[:5]:
        summary_lines.append(f"  - {font}: {count} блоків")

    summary_lines.append(f"\nРозміри шрифтів:")
    for size, count in sorted(font_sizes.items(), key=lambda x: -x[1])[:5]:
        summary_lines.append(f"  - {size}pt: {count} блоків")

    # Виявлені розділи документа
    sections = []
    seen_sections = set()
    for block in blocks:
        section = block.get("section", "")
        if section and section not in seen_sections:
            seen_sections.add(section)
            sections.append(section)

    if sections:
        summary_lines.append(f"\nСТРУКТУРА ДОКУМЕНТА (розділи):")
        for i, section in enumerate(sections[:20], 1):  # Перші 20 розділів
            summary_lines.append(f"  {i}. {section}")

    # Приклади блоків з різним форматуванням
    summary_lines.append(f"\nПРИКЛАДИ БЛОКІВ З ФОРМАТУВАННЯМ:\n")

    # Знаходимо унікальні комбінації форматування
    seen_combinations = set()
    examples = []

    for block in blocks[:50]:  # Перевіряємо перші 50 блоків
        combination = (block.get("alignment"), block.get("font_size"), block.get("bold"))
        if combination not in seen_combinations and len(examples) < 10:
            seen_combinations.add(combination)
            examples.append(block)

    for i, block in enumerate(examples, 1):
        summary_lines.append(
            f"{i}. [РОЗДІЛ: {block.get('section', 'N/A')[:40]}] "
            f"[{block.get('alignment', 'left').upper()}, "
            f"{block.get('font_name', 'Unknown')}, "
            f"{block.get('font_size', 0)}pt, "
            f"{'BOLD' if block.get('bold') else 'normal'}] "
            f'"{block.get("text", "")[:60]}..."'
        )

    return "\n".join(summary_lines)


def format_blocks_for_validation(blocks: List[Dict], start_idx: int = 0, limit: int = 50) -> str:
    """Форматує блоки для валідації з інформацією про розділи та форматування

    Args:
        blocks: Список блоків з форматуванням
        start_idx: З якого блоку починати
        limit: Скільки блоків включити

    Returns:
        Текст для передачі в LLM валідатор
    """
    selected_blocks = blocks[start_idx:start_idx + limit]
    if not selected_blocks:
        return ""

    lines = []
    current_section = None

    for block in selected_blocks:
        section = block.get("section", "")

        # Якщо змінився розділ - додаємо заголовок
        if section != current_section:
            lines.append(f"\n--- {section} ---\n")
            current_section = section

        # Форматуємо блок з метаданими
        alignment = block.get("alignment", "left")
        font_name = block.get("font_name", "Unknown")
        font_size = block.get("font_size", 0)
        bold = "BOLD" if block.get("bold") else "normal"

        lines.append(
            f"[{alignment.upper()}|{font_name}|{font_size}pt|{bold}] {block.get('text', '')}"
        )

    return "\n".join(lines)