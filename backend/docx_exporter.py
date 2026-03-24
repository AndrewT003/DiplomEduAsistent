from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from typing import Dict

def set_times_new_roman_14pt(run):
    """Встановлює Times New Roman 14pt і чорний колір для тексту"""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Чорний колір
    # Також встановлюємо для складних скриптів
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def apply_font_to_document(doc):
    """Застосовує Times New Roman 14pt та вирівнювання до всього документа"""
    for paragraph in doc.paragraphs:
        # Визначаємо чи це заголовок
        is_heading = paragraph.style.name.startswith('Heading') or paragraph.style.name == 'Title'

        # Встановлюємо вирівнювання
        if is_heading or paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            # Зберігаємо оригінальний колір якщо він не чорний (для кольорових елементів)
            if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0, 0, 0):
                run.font.color.rgb = RGBColor(0, 0, 0)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def set_paragraph_alignment(paragraph, is_heading=False):
    """Встановлює вирівнювання параграфа: заголовки - по центру, текст - по ширині"""
    if is_heading:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_text_with_inline_markdown(paragraph, text):
    """Додає текст до параграфа з обробкою inline markdown (**жирний**, *курсив*)"""
    import re

    # Простий парсинг: спочатку обробляємо жирний текст **
    parts = []
    bold_pattern = r'\*\*(.+?)\*\*'
    italic_pattern = r'(?<!\*)\*(.+?)\*(?!\*)'

    # Розбиваємо текст на частини
    current_text = text
    last_pos = 0
    result_parts = []

    # Обробляємо жирний текст
    for match in re.finditer(bold_pattern, text):
        # Додаємо текст до match
        if match.start() > last_pos:
            result_parts.append(('normal', text[last_pos:match.start()]))

        # Додаємо жирний текст
        result_parts.append(('bold', match.group(1)))
        last_pos = match.end()

    # Додаємо решту тексту
    if last_pos < len(text):
        result_parts.append(('normal', text[last_pos:]))

    # Тепер обробляємо курсив у нормальних частинах
    final_parts = []
    for part_type, part_text in result_parts:
        if part_type == 'normal':
            # Обробляємо курсив
            last_italic_pos = 0
            for match in re.finditer(italic_pattern, part_text):
                if match.start() > last_italic_pos:
                    final_parts.append(('normal', part_text[last_italic_pos:match.start()]))

                final_parts.append(('italic', match.group(1)))
                last_italic_pos = match.end()

            if last_italic_pos < len(part_text):
                final_parts.append(('normal', part_text[last_italic_pos:]))
        else:
            final_parts.append((part_type, part_text))

    # Додаємо всі частини до параграфа
    for part_type, part_text in final_parts:
        if not part_text:
            continue

        run = paragraph.add_run(part_text)
        if part_type == 'bold':
            run.bold = True
        elif part_type == 'italic':
            run.italic = True

def export_to_docx(content: str, title: str, material_type: str, topic: str = None) -> bytes:
    doc = Document()

    # Заголовок документу - використовуємо тему якщо є
    display_title = topic if topic else title
    heading = doc.add_heading(display_title, level=0)
    set_paragraph_alignment(heading, is_heading=True)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Тип матеріалу як підзаголовок
    labels = {
        "summary":    "Конспект",
        "quiz":       "Тестові питання",
        "flashcards": "Флеш-картки",
        "glossary":   "Глосарій"
    }
    sub = doc.add_paragraph(labels.get(material_type, "Матеріал"))
    set_paragraph_alignment(sub, is_heading=True)
    for run in sub.runs:
        set_times_new_roman_14pt(run)

    doc.add_paragraph()

    # Парсимо контент по рядках
    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue

        # Перевіряємо чи це заголовок з markdown
        if line.startswith("##### "):
            h = doc.add_heading(line[6:], level=5)
            set_paragraph_alignment(h, is_heading=True)
            for run in h.runs:
                set_times_new_roman_14pt(run)
        elif line.startswith("#### "):
            h = doc.add_heading(line[5:], level=4)
            set_paragraph_alignment(h, is_heading=True)
            for run in h.runs:
                set_times_new_roman_14pt(run)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:], level=3)
            set_paragraph_alignment(h, is_heading=True)
            for run in h.runs:
                set_times_new_roman_14pt(run)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:], level=2)
            set_paragraph_alignment(h, is_heading=True)
            for run in h.runs:
                set_times_new_roman_14pt(run)
        elif line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            set_paragraph_alignment(h, is_heading=True)
            for run in h.runs:
                set_times_new_roman_14pt(run)
        elif line.startswith("- ") or line.startswith("* "):
            # Прибираємо маркер списку
            text = line[2:]
            # Обробляємо inline markdown
            p = doc.add_paragraph(style="List Bullet")
            set_paragraph_alignment(p, is_heading=False)
            add_text_with_inline_markdown(p, text)
            for run in p.runs:
                set_times_new_roman_14pt(run)
        else:
            # Звичайний параграф - обробляємо inline markdown
            p = doc.add_paragraph()
            set_paragraph_alignment(p, is_heading=False)
            add_text_with_inline_markdown(p, line)
            for run in p.runs:
                set_times_new_roman_14pt(run)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def export_validation_to_docx(validation_result: Dict, document_name: str = "Документ") -> bytes:
    """Експортує звіт валідації у форматований DOCX"""
    doc = Document()

    # Заголовок
    heading = doc.add_heading("Звіт валідації документа", level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    # Назва документа
    doc_para = doc.add_paragraph(f"Документ: {document_name}")
    doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in doc_para.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.italic = True
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph()

    # Загальна інформація
    compliance = validation_result.get("overall_compliance", "unknown")
    score = validation_result.get("compliance_score", 0.0)
    summary = validation_result.get("summary", {})

    # Статус відповідності
    status_para = doc.add_paragraph()
    status_para.add_run("Статус відповідності: ").bold = True

    status_run = status_para.add_run(compliance.upper())
    status_run.bold = True

    # Колір залежно від статусу
    if compliance == "pass":
        status_run.font.color.rgb = RGBColor(0, 128, 0)  # Зелений
    elif compliance == "partial":
        status_run.font.color.rgb = RGBColor(255, 165, 0)  # Помаранчевий
    elif compliance == "fail":
        status_run.font.color.rgb = RGBColor(255, 0, 0)  # Червоний

    # Оцінка
    score_para = doc.add_paragraph()
    score_para.add_run("Оцінка відповідності: ").bold = True
    score_para.add_run(f"{score * 100:.1f}%")

    doc.add_paragraph()

    # Підсумок проблем
    doc.add_heading("Підсумок проблем", level=1)

    total = summary.get("total_issues", 0)
    if total == 0:
        success_para = doc.add_paragraph("Проблем не виявлено! Документ відповідає всім нормативним вимогам.")
        success_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)
        success_para.runs[0].bold = True
    else:
        doc.add_paragraph(f"Загальна кількість проблем: {total}", style="List Bullet")
        doc.add_paragraph(f"Критичні: {summary.get('critical', 0)}", style="List Bullet 2")
        doc.add_paragraph(f"Важливі: {summary.get('major', 0)}", style="List Bullet 2")
        doc.add_paragraph(f"Незначні: {summary.get('minor', 0)}", style="List Bullet 2")
        doc.add_paragraph(f"Інформаційні: {summary.get('info', 0)}", style="List Bullet 2")

    doc.add_paragraph()

    # Детальний список проблем
    issues = validation_result.get("issues", [])

    if issues:
        doc.add_heading("Виявлені проблеми", level=1)

        # Групуємо по severity
        severity_labels = {
            "critical": "Критичні проблеми",
            "major": "Важливі проблеми",
            "minor": "Незначні проблеми",
            "info": "Інформаційні зауваження"
        }

        severity_colors = {
            "critical": RGBColor(255, 0, 0),
            "major": RGBColor(255, 165, 0),
            "minor": RGBColor(255, 215, 0),
            "info": RGBColor(0, 0, 255)
        }

        for severity in ["critical", "major", "minor", "info"]:
            severity_issues = [i for i in issues if i.get("severity") == severity]

            if severity_issues:
                severity_heading = doc.add_heading(severity_labels[severity], level=2)
                severity_heading.runs[0].font.color.rgb = severity_colors[severity]

                for idx, issue in enumerate(severity_issues, 1):
                    # Опис проблеми
                    issue_para = doc.add_paragraph()
                    issue_num = issue_para.add_run(f"{idx}. ")
                    issue_num.bold = True
                    issue_desc = issue_para.add_run(issue.get("description", "Невідома проблема"))
                    issue_desc.bold = True

                    # Нормативний документ
                    doc.add_paragraph(
                        f"Нормативний документ: {issue.get('regulatory_doc_name', 'N/A')}",
                        style="List Bullet"
                    )

                    # Розділ
                    doc.add_paragraph(
                        f"Розділ нормативу: {issue.get('section', 'N/A')}",
                        style="List Bullet"
                    )

                    # Категорія
                    doc.add_paragraph(
                        f"Категорія: {issue.get('category', 'N/A')}",
                        style="List Bullet"
                    )

                    # Фрагмент документа
                    fragment = issue.get("user_doc_fragment", "")
                    if fragment:
                        frag_para = doc.add_paragraph(style="List Bullet")
                        frag_para.add_run("Фрагмент документа: ").italic = True
                        frag_para.add_run(f'"{fragment[:200]}"')

                    # Рекомендація
                    recommendation = issue.get("recommendation", "")
                    if recommendation:
                        rec_para = doc.add_paragraph(style="List Bullet")
                        rec_para.add_run("Рекомендація: ").bold = True
                        rec_para.add_run(recommendation)

                    doc.add_paragraph()  # Пустий рядок між проблемами

    # Загальні рекомендації
    doc.add_heading("Загальні рекомендації", level=1)

    if compliance == "pass":
        rec_text = "Документ повністю відповідає нормативним вимогам. Чудова робота!"
    elif compliance == "partial":
        rec_text = "Документ частково відповідає вимогам. Рекомендується виправити виявлені проблеми для повної відповідності."
    elif compliance == "fail":
        rec_text = "Документ має значні невідповідності. Необхідно терміново виправити критичні та важливі проблеми перед поданням."
    else:
        rec_text = "Не вдалося визначити рівень відповідності. Перевірте звіт детально."

    doc.add_paragraph(rec_text)

    # Застосовуємо Times New Roman 14pt до всього документа
    apply_font_to_document(doc)

    # Збереження у BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()