"""
Спеціалізовані валідатори для перевірки форматування документів
Використовують python-docx для точного аналізу (без LLM)
"""

import re
from typing import List, Dict, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class FormattingValidator:
    """Валідатор форматування для DOCX документів"""

    def __init__(self, file_path: str):
        """
        Args:
            file_path: Шлях до DOCX файлу
        """
        try:
            self.doc = Document(file_path)
            self.file_path = file_path
        except Exception as e:
            raise ValueError(f"Не вдалося відкрити документ: {e}")

    def validate_font(
        self,
        expected_font: str = "Times New Roman",
        expected_size: int = 14,
        tolerance: float = 0.5
    ) -> List[Dict]:
        """Перевірка шрифту основного тексту

        Args:
            expected_font: Очікувана назва шрифту
            expected_size: Очікуваний розмір в пунктах
            tolerance: Допустиме відхилення розміру

        Returns:
            Список проблем з форматуванням
        """
        issues = []
        paragraph_count = 0
        wrong_font_count = 0
        wrong_size_count = 0

        for para_idx, para in enumerate(self.doc.paragraphs):
            # Пропускаємо порожні параграфи
            if not para.text.strip():
                continue

            paragraph_count += 1

            if not para.runs:
                continue

            # Аналізуємо перший run (зазвичай він визначає стиль параграфа)
            run = para.runs[0]

            # Перевірка назви шрифту
            if run.font.name and run.font.name != expected_font:
                wrong_font_count += 1

                # Додаємо issue тільки для перших 3 випадків (щоб не засмічувати звіт)
                if wrong_font_count <= 3:
                    issues.append({
                        "severity": "major",
                        "category": "formatting",
                        "description": f"Неправильний шрифт: '{run.font.name}' (очікувався '{expected_font}')",
                        "user_document_section": self._get_section_name(para_idx),
                        "user_doc_fragment": para.text[:150],
                        "recommendation": f"Змініть шрифт на '{expected_font}' в усьому документі"
                    })

            # Перевірка розміру шрифту
            if run.font.size:
                size_pt = run.font.size.pt
                if abs(size_pt - expected_size) > tolerance:
                    wrong_size_count += 1

                    if wrong_size_count <= 3:
                        issues.append({
                            "severity": "minor",
                            "category": "formatting",
                            "description": f"Неправильний розмір шрифту: {size_pt}pt (очікувався {expected_size}pt)",
                            "user_document_section": self._get_section_name(para_idx),
                            "user_doc_fragment": para.text[:150],
                            "recommendation": f"Встановіть розмір шрифту {expected_size}pt для основного тексту"
                        })

        # Якщо більше 50% параграфів мають неправильне форматування - додаємо загальне попередження
        if paragraph_count > 0:
            if wrong_font_count > paragraph_count * 0.5:
                issues.insert(0, {
                    "severity": "critical",
                    "category": "formatting",
                    "description": f"Більшість тексту ({wrong_font_count}/{paragraph_count} параграфів) має неправильний шрифт",
                    "user_document_section": "ВЕСЬ ДОКУМЕНТ",
                    "user_doc_fragment": "",
                    "recommendation": f"Виділіть весь текст (Ctrl+A) та встановіть шрифт '{expected_font}'"
                })

            if wrong_size_count > paragraph_count * 0.5:
                issues.insert(0, {
                    "severity": "major",
                    "category": "formatting",
                    "description": f"Більшість тексту ({wrong_size_count}/{paragraph_count} параграфів) має неправильний розмір шрифту",
                    "user_document_section": "ВЕСЬ ДОКУМЕНТ",
                    "user_doc_fragment": "",
                    "recommendation": f"Виділіть весь текст (Ctrl+A) та встановіть розмір {expected_size}pt"
                })

        return issues

    def validate_line_spacing(
        self,
        expected_spacing: float = 1.5,
        tolerance: float = 0.1
    ) -> List[Dict]:
        """Перевірка міжрядкового інтервалу

        Args:
            expected_spacing: Очікуваний інтервал (1.0, 1.5, 2.0 тощо)
            tolerance: Допустиме відхилення

        Returns:
            Список проблем з інтервалами
        """
        issues = []
        wrong_spacing_count = 0

        for para_idx, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue

            # Отримуємо міжрядковий інтервал
            spacing = para.paragraph_format.line_spacing

            if spacing and isinstance(spacing, float):
                if abs(spacing - expected_spacing) > tolerance:
                    wrong_spacing_count += 1

                    if wrong_spacing_count <= 3:
                        issues.append({
                            "severity": "minor",
                            "category": "formatting",
                            "description": f"Неправильний міжрядковий інтервал: {spacing:.1f} (очікувався {expected_spacing})",
                            "user_document_section": self._get_section_name(para_idx),
                            "user_doc_fragment": para.text[:150],
                            "recommendation": f"Встановіть міжрядковий інтервал {expected_spacing}"
                        })

        return issues

    def validate_paragraph_indent(
        self,
        expected_indent_cm: float = 1.25,
        tolerance_cm: float = 0.1
    ) -> List[Dict]:
        """Перевірка абзацного відступу

        Args:
            expected_indent_cm: Очікуваний відступ в сантиметрах
            tolerance_cm: Допустиме відхилення

        Returns:
            Список проблем з відступами
        """
        issues = []
        wrong_indent_count = 0

        for para_idx, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue

            # Пропускаємо заголовки (зазвичай без відступу)
            if para.style.name.startswith('Heading'):
                continue

            indent = para.paragraph_format.first_line_indent

            if indent is not None:
                indent_cm = indent.cm

                if abs(indent_cm - expected_indent_cm) > tolerance_cm:
                    wrong_indent_count += 1

                    if wrong_indent_count <= 3:
                        issues.append({
                            "severity": "minor",
                            "category": "formatting",
                            "description": f"Неправильний абзацний відступ: {indent_cm:.2f}см (очікувався {expected_indent_cm:.2f}см)",
                            "user_document_section": self._get_section_name(para_idx),
                            "user_doc_fragment": para.text[:150],
                            "recommendation": f"Встановіть абзацний відступ {expected_indent_cm}см (Формат → Абзац → Перший рядок → Відступ)"
                        })
            else:
                # Відступ не встановлено
                wrong_indent_count += 1

                if wrong_indent_count <= 3:
                    issues.append({
                        "severity": "info",
                        "category": "formatting",
                        "description": "Абзацний відступ не встановлено",
                        "user_document_section": self._get_section_name(para_idx),
                        "user_doc_fragment": para.text[:150],
                        "recommendation": f"Встановіть абзацний відступ {expected_indent_cm}см"
                    })

        return issues

    def validate_margins(
        self,
        top_cm: float = 2.0,
        bottom_cm: float = 2.0,
        left_cm: float = 3.0,
        right_cm: float = 1.5,
        tolerance_cm: float = 0.2
    ) -> List[Dict]:
        """Перевірка полів сторінки

        Args:
            top_cm: Верхнє поле (см)
            bottom_cm: Нижнє поле (см)
            left_cm: Ліве поле (см)
            right_cm: Праве поле (см)
            tolerance_cm: Допустиме відхилення

        Returns:
            Список проблем з полями
        """
        issues = []

        # Отримуємо налаштування розділу (зазвичай один розділ на документ)
        for section in self.doc.sections:
            actual_top = section.top_margin.cm
            actual_bottom = section.bottom_margin.cm
            actual_left = section.left_margin.cm
            actual_right = section.right_margin.cm

            margins_check = [
                ("верхнє", actual_top, top_cm),
                ("нижнє", actual_bottom, bottom_cm),
                ("ліве", actual_left, left_cm),
                ("праве", actual_right, right_cm)
            ]

            for margin_name, actual, expected in margins_check:
                if abs(actual - expected) > tolerance_cm:
                    issues.append({
                        "severity": "major",
                        "category": "formatting",
                        "description": f"Неправильне {margin_name} поле: {actual:.1f}см (очікувалось {expected:.1f}см)",
                        "user_document_section": "НАЛАШТУВАННЯ СТОРІНКИ",
                        "user_doc_fragment": "",
                        "recommendation": f"Встановіть {margin_name} поле {expected:.1f}см (Макет → Поля → Настроюване)"
                    })

        return issues

    def validate_page_size(
        self,
        expected_width_cm: float = 21.0,  # A4
        expected_height_cm: float = 29.7,  # A4
        tolerance_cm: float = 0.5
    ) -> List[Dict]:
        """Перевірка розміру сторінки

        Args:
            expected_width_cm: Очікувана ширина (см)
            expected_height_cm: Очікувана висота (см)
            tolerance_cm: Допустиме відхилення

        Returns:
            Список проблем з розміром сторінки
        """
        issues = []

        for section in self.doc.sections:
            actual_width = section.page_width.cm
            actual_height = section.page_height.cm

            if abs(actual_width - expected_width_cm) > tolerance_cm or \
               abs(actual_height - expected_height_cm) > tolerance_cm:
                issues.append({
                    "severity": "major",
                    "category": "formatting",
                    "description": f"Неправильний розмір сторінки: {actual_width:.1f}×{actual_height:.1f}см (очікувався A4: {expected_width_cm}×{expected_height_cm}см)",
                    "user_document_section": "НАЛАШТУВАННЯ СТОРІНКИ",
                    "user_doc_fragment": "",
                    "recommendation": "Встановіть розмір сторінки A4 (Макет → Розмір → A4)"
                })

        return issues

    def validate_alignment(
        self,
        expected_alignment: str = "justify"  # left, center, right, justify
    ) -> List[Dict]:
        """Перевірка вирівнювання тексту

        Args:
            expected_alignment: Очікуване вирівнювання

        Returns:
            Список проблем з вирівнюванням
        """
        issues = []
        wrong_alignment_count = 0

        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
        }

        for para_idx, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():
                continue

            # Пропускаємо заголовки (вони можуть мати інше вирівнювання)
            if para.style.name.startswith('Heading'):
                continue

            actual_alignment = alignment_map.get(para.alignment, "left")

            if actual_alignment != expected_alignment:
                wrong_alignment_count += 1

                if wrong_alignment_count <= 3:
                    alignment_names = {
                        "left": "по лівому краю",
                        "center": "по центру",
                        "right": "по правому краю",
                        "justify": "по ширині"
                    }

                    issues.append({
                        "severity": "minor",
                        "category": "formatting",
                        "description": f"Неправильне вирівнювання: {alignment_names[actual_alignment]} (очікувалось {alignment_names[expected_alignment]})",
                        "user_document_section": self._get_section_name(para_idx),
                        "user_doc_fragment": para.text[:150],
                        "recommendation": f"Встановіть вирівнювання {alignment_names[expected_alignment]}"
                    })

        return issues

    def validate_all(
        self,
        rules: List[Dict] = None
    ) -> List[Dict]:
        """Виконати всі перевірки форматування

        Args:
            rules: Список правил форматування з нормативного документа (опціонально)
                   Якщо не вказано - використовуються стандартні ДСТУ

        Returns:
            Загальний список проблем
        """
        all_issues = []

        print("  🎨 Перевірка форматування (без LLM)...")

        # Стандартні перевірки ДСТУ
        all_issues.extend(self.validate_font(
            expected_font="Times New Roman",
            expected_size=14
        ))

        all_issues.extend(self.validate_line_spacing(
            expected_spacing=1.5
        ))

        all_issues.extend(self.validate_paragraph_indent(
            expected_indent_cm=1.25
        ))

        all_issues.extend(self.validate_margins(
            top_cm=2.0,
            bottom_cm=2.0,
            left_cm=3.0,
            right_cm=1.5
        ))

        all_issues.extend(self.validate_page_size(
            expected_width_cm=21.0,
            expected_height_cm=29.7
        ))

        all_issues.extend(self.validate_alignment(
            expected_alignment="justify"
        ))

        print(f"  ✅ Перевірка форматування: знайдено {len(all_issues)} проблем")

        return all_issues

    def _get_section_name(self, para_idx: int) -> str:
        """Визначити назву розділу для параграфа"""
        # Простий підхід: шукаємо найближчий заголовок перед цим параграфом
        for i in range(para_idx, -1, -1):
            para = self.doc.paragraphs[i]
            if para.style.name.startswith('Heading'):
                return para.text.strip()[:100]

        return "ПОЧАТОК ДОКУМЕНТА"


class StructureValidator:
    """Валідатор структури документа"""

    def __init__(self, chunks: List[Dict]):
        """
        Args:
            chunks: Чанки документа з інформацією про розділи
        """
        self.chunks = chunks
        self.sections = self._extract_sections()

    def _extract_sections(self) -> List[str]:
        """Витягти всі розділи з документа"""
        sections = []
        seen = set()

        for chunk in self.chunks:
            section = chunk.get("section", "")
            if section and section not in seen:
                sections.append(section)
                seen.add(section)

        return sections

    def validate_required_sections(
        self,
        required: List[str]
    ) -> List[Dict]:
        """Перевірка наявності обов'язкових розділів

        Args:
            required: Список обов'язкових розділів

        Returns:
            Список проблем з структурою
        """
        issues = []

        print(f"  📚 Перевірка структури: {len(required)} обов'язкових розділів...")

        for req_section in required:
            # Нечіткий пошук (враховуємо варіації написання)
            found = any(
                req_section.lower() in section.lower()
                for section in self.sections
            )

            if not found:
                # Визначаємо severity в залежності від типу розділу
                severity = "critical" if any(
                    keyword in req_section.lower()
                    for keyword in ["вступ", "висновки", "список"]
                ) else "major"

                issues.append({
                    "severity": severity,
                    "category": "structure",
                    "description": f"Відсутній обов'язковий розділ: '{req_section}'",
                    "user_document_section": "СТРУКТУРА ДОКУМЕНТА",
                    "user_doc_fragment": "",
                    "recommendation": f"Додайте розділ '{req_section}' до документа"
                })

        print(f"  ✅ Перевірка структури: знайдено {len(issues)} проблем")
        return issues

    def validate_section_order(
        self,
        expected_order: List[str]
    ) -> List[Dict]:
        """Перевірка порядку розділів

        Args:
            expected_order: Очікуваний порядок розділів

        Returns:
            Список проблем з порядком
        """
        issues = []

        # Знаходимо позиції розділів
        section_positions = {}
        for idx, section in enumerate(self.sections):
            for expected in expected_order:
                if expected.lower() in section.lower():
                    section_positions[expected] = idx
                    break

        # Перевіряємо порядок
        prev_position = -1
        for expected in expected_order:
            if expected in section_positions:
                position = section_positions[expected]
                if position < prev_position:
                    issues.append({
                        "severity": "major",
                        "category": "structure",
                        "description": f"Розділ '{expected}' знаходиться не на своєму місці в документі",
                        "user_document_section": "СТРУКТУРА ДОКУМЕНТА",
                        "user_doc_fragment": "",
                        "recommendation": f"Перемістіть розділ '{expected}' у правильну позицію"
                    })
                prev_position = position

        return issues


class ReferencesValidator:
    """Валідатор посилань та цитувань"""

    def __init__(self, text: str):
        """
        Args:
            text: Повний текст документа
        """
        self.text = text

    def validate_citations(self) -> List[Dict]:
        """Перевірка посилань [1], [2] тощо"""
        issues = []

        print(f"  🔗 Перевірка посилань...")

        # Знаходимо всі посилання [1], [2], etc.
        citations = re.findall(r'\[(\d+)\]', self.text)

        if not citations:
            print(f"  ℹ️ Посилань не знайдено")
            return []

        # Перевіряємо чи є розділ "Список використаних джерел"
        has_references_section = any(
            keyword in self.text.lower()
            for keyword in [
                "список використаних джерел",
                "список літератури",
                "бібліографія",
                "references"
            ]
        )

        if not has_references_section:
            issues.append({
                "severity": "critical",
                "category": "references",
                "description": f"Знайдено {len(citations)} посилань в тексті, але відсутній розділ 'Список використаних джерел'",
                "user_document_section": "СТРУКТУРА ДОКУМЕНТА",
                "user_doc_fragment": "",
                "recommendation": "Додайте розділ 'Список використаних джерел' в кінці документа"
            })

        # Перевіряємо послідовність нумерації
        citations_nums = [int(c) for c in citations]
        unique_citations = sorted(set(citations_nums))

        for i, num in enumerate(unique_citations, 1):
            if num != i:
                issues.append({
                    "severity": "major",
                    "category": "references",
                    "description": f"Порушення нумерації посилань: знайдено [{num}], очікувалось [{i}]",
                    "user_document_section": "ПОСИЛАННЯ",
                    "user_doc_fragment": "",
                    "recommendation": "Перенумеруйте посилання послідовно: [1], [2], [3], ..."
                })
                break

        # Перевіряємо чи всі посилання унікальні місця (не [1][1][1])
        if len(citations) > len(unique_citations) * 3:
            issues.append({
                "severity": "info",
                "category": "references",
                "description": f"Велика кількість повторень посилань: {len(citations)} посилань на {len(unique_citations)} джерел",
                "user_document_section": "ПОСИЛАННЯ",
                "user_doc_fragment": "",
                "recommendation": "Перевірте чи всі повторні посилання коректні"
            })

        print(f"  ✅ Перевірка посилань: знайдено {len(issues)} проблем")
        return issues
